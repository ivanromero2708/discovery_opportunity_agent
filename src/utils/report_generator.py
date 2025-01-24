# src/utils/report_generator.py

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_markdown_to_word(report: str, industry: str, region: str) -> str:
    """
    Convierte un reporte en formato Markdown a un documento Word con formato profesional.
    
    Args:
        report (str): Contenido del reporte en Markdown.
        industry (str): Industria objetivo.
        region (str): Región objetivo.
    
    Returns:
        str: Ruta del archivo generado.
    """
    try:
        # Nos aseguramos de que exista la carpeta "reports"
        os.makedirs("reports", exist_ok=True)
        
        doc = Document()
        
        # Ajuste de estilos: ejemplo de estilo 'Heading 1'
        styles = doc.styles
        heading_style = styles['Heading 1']
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(16)
        heading_font.bold = True

        # Metadatos del documento
        doc.core_properties.title = f"Market Research Report - {industry} - {region}"
        doc.core_properties.author = "AI Research Assistant"

        # Procesar contenido (dividir por líneas, buscar #, ##, etc.)
        lines = report.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                # Texto normal
                para = doc.add_paragraph(line)
                para_format = para.paragraph_format
                para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Generar nombre de archivo
        filename = f"reports/Market_Research_{industry.replace(' ', '_')}_{region.replace(' ', '_')}.docx"
        doc.save(filename)
        
        return filename
        
    except Exception as e:
        raise RuntimeError(f"Error en la generación de Word: {str(e)}")
